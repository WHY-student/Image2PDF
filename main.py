import os
import shutil
from docx import Document
from docx.shared import Inches
import win32com.client as win32
from PIL import Image

def resize_image(input_image, output_image, new_width):
    image = Image.open(input_image)
    width, height = image.size
    aspect_ratio = height / width
    new_height = int(new_width * aspect_ratio)
    resized_image = image.resize((new_width, new_height), Image.ANTIALIAS)
    resized_image.save(output_image)

def resize_image1(image_path):
    new_width = 430  # 新的图片宽度
    resize_image(image_path, image_path, new_width)

def insert_images_to_word(image_folder, output_docx):
    image_folder = "./test/"+image_folder
    doc = Document()
    for filename in os.listdir(image_folder):
        if filename.endswith(".jpg") or filename.endswith(".png"):
            image_path = os.path.join(image_folder, filename)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6.0))
    doc.save(output_docx)

def convert_to_pdf(input_docx, output_pdf):
    word = win32.Dispatch("Word.Application")
    doc = word.Documents.Open("F:/情若炎兮/Desktop/do_everything/"+input_docx)
    doc.SaveAs2("F:/情若炎兮/Desktop/do_everything/pdf1/"+output_pdf, FileFormat=17)
    doc.Close()
    word.Quit()

def get_subdirectories(folder_path):
    subdirectories = []
    for item in os.listdir(folder_path):
        item_path = os.path.join(folder_path, item)
        if os.path.isdir(item_path):
            subdirectories.append(item)
    return subdirectories

def rename_image(folder_path):
    # 获取文件夹下的所有文件
    files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

    # 过滤出图片文件
    image_files = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
    print(image_files)
    # 按照修改时间进行排序
    sorted_files = sorted(image_files, key=lambda f: os.path.getctime(os.path.join(folder_path, f)))
    print(sorted_files)
    # 遍历排序后的文件列表，进行重命名操作
    for i, filename in enumerate(sorted_files):
        file_path = os.path.join(folder_path, filename)
        new_filename = '{}.{}'.format(i + 1, filename.split('.')[-1])
        new_file_path = os.path.join(folder_path, new_filename)
        shutil.move(file_path, new_file_path)



def main():
    root_dir = "F:/情若炎兮/Desktop/do_everything/test"
    subdirectories = get_subdirectories(root_dir)
    print("Subdirectories:")
    for subdirectory in subdirectories:
    # name = "张金达"
        name = subdirectory
        print(name)
        image_folder = name+"/"  # 替换为图片所在的文件夹路径
        output_docx = name+".docx"  # 输出的Word文档路径
        output_pdf = name+"入团志愿书.pdf"  # 输出的PDF文件路径
        # if name in ["吴媛媛", "吴梦娴"]:
        rename_image(root_dir+"/"+name)
        insert_images_to_word(image_folder, output_docx)
        convert_to_pdf(output_docx, output_pdf)

if __name__ == "__main__":
    main()



